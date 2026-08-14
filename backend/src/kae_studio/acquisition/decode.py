"""Turn uploaded bytes into text Memory can ingest (`VC-06/H`).

Memory's document route is **text only**. Decoding is Studio's problem, and
doing it here — on the trusted backend, in process, with no third-party
reader — keeps briefs and spreadsheets on this host.

This is not analysis. The file becomes evidence; extraction proposes
statements; a person confirms. A scanned PDF that yields almost no text is a
warning, not a silent empty success, and types we cannot read are refused
rather than accepted and lost.
"""

from __future__ import annotations

import csv
import io
from dataclasses import dataclass, field

MAX_BYTES = 10 * 1024 * 1024
MAX_EXCEL_ROWS = 200
MAX_EXCEL_SHEETS = 8
# Digital PDFs that are actually scans produce almost nothing. Below this,
# say so rather than ingesting an empty document that looks like it was read.
MIN_PDF_CHARS = 40

SUPPORTED = (
    "PDF, Word (.docx), Excel (.xlsx), CSV, and plain text or Markdown. "
    "Legacy .doc, scanned PDFs (OCR), images, email and zip are not read yet."
)


class DecodeError(ValueError):
    """The file cannot be turned into text, in one sentence a person can act on."""


@dataclass(frozen=True)
class DecodedDocument:
    text: str
    warnings: list[str] = field(default_factory=list)
    format: str = ""
    suggested_title: str = ""


def decode(data: bytes, filename: str, content_type: str = "") -> DecodedDocument:
    """Extract text from an allowed upload, or raise `DecodeError`."""

    if len(data) > MAX_BYTES:
        raise DecodeError(
            f"This file is larger than {MAX_BYTES // (1024 * 1024)} MB, which KAE will not read. "
            "Paste the relevant passages, or split the file."
        )
    if not data:
        raise DecodeError("The file is empty, so there is nothing to read.")

    kind = _kind(filename, content_type)
    name = filename.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    title = name.rsplit(".", 1)[0] if "." in name else name

    if kind == "pdf":
        text, warnings = _pdf(data)
    elif kind == "docx":
        text, warnings = _docx(data)
    elif kind == "xlsx":
        text, warnings = _xlsx(data)
    elif kind == "csv":
        text, warnings = _csv(data)
    elif kind == "text":
        text, warnings = _text(data)
    else:
        raise DecodeError(
            f"KAE cannot read this file type. {SUPPORTED} Save as one of those, or paste the text."
        )

    stripped = text.strip()
    if not stripped:
        raise DecodeError(
            "KAE found no readable text in this file. A scanned PDF needs OCR, which is not "
            "available, and an empty workbook has nothing to extract. Paste the text if you have it."
        )
    return DecodedDocument(text=stripped, warnings=warnings, format=kind, suggested_title=title)


def _kind(filename: str, content_type: str) -> str:
    name = filename.lower()
    mime = (content_type or "").split(";", 1)[0].strip().lower()
    if name.endswith(".pdf") or mime == "application/pdf":
        return "pdf"
    if name.endswith(".docx") or mime in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }:
        return "docx"
    if name.endswith(".xlsx") or mime in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }:
        return "xlsx"
    if name.endswith(".csv") or mime == "text/csv":
        return "csv"
    if name.endswith((".txt", ".md", ".markdown")) or mime in {"text/plain", "text/markdown"}:
        return "text"
    if name.endswith(".doc") or mime == "application/msword":
        raise DecodeError(
            "Legacy Word (.doc) is not read. Save the file as .docx, or paste the text."
        )
    return ""


def _pdf(data: bytes) -> tuple[str, list[str]]:
    try:
        import pymupdf
    except ImportError as error:
        raise DecodeError("PDF reading is not installed on this Studio.") from error

    try:
        document = pymupdf.open(stream=data, filetype="pdf")
    except Exception as error:
        raise DecodeError(
            "This file could not be read as a PDF. It may be damaged. Paste the text if you have it."
        ) from error
    try:
        pages = [page.get_text("text") or "" for page in document]
        page_count = document.page_count
    finally:
        document.close()

    text = "\n\n".join(part.strip() for part in pages if part.strip())
    warnings: list[str] = []
    if page_count and len(text) < MIN_PDF_CHARS:
        warnings.append(
            "This PDF has almost no extractable text. It may be a scan. OCR is not available; "
            "paste the wording if you have it."
        )
    return text, warnings


def _docx(data: bytes) -> tuple[str, list[str]]:
    try:
        from docx import Document
    except ImportError as error:
        raise DecodeError("Word reading is not installed on this Studio.") from error

    try:
        document = Document(io.BytesIO(data))
    except Exception as error:
        raise DecodeError(
            "This file could not be read as Word. Save it as .docx, or paste the text."
        ) from error
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts), []


def _xlsx(data: bytes) -> tuple[str, list[str]]:
    try:
        from openpyxl import load_workbook
    except ImportError as error:
        raise DecodeError("Excel reading is not installed on this Studio.") from error

    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as error:
        raise DecodeError(
            "This file could not be read as Excel. Save it as .xlsx, or paste the relevant rows."
        ) from error
    warnings: list[str] = []
    sheets = workbook.sheetnames
    if len(sheets) > MAX_EXCEL_SHEETS:
        warnings.append(
            f"Only the first {MAX_EXCEL_SHEETS} sheets were read; this workbook has {len(sheets)}."
        )
        sheets = sheets[:MAX_EXCEL_SHEETS]

    blocks: list[str] = []
    try:
        for name in sheets:
            rows_out: list[str] = []
            truncated = False
            for index, row in enumerate(workbook[name].iter_rows(values_only=True), start=1):
                if index > MAX_EXCEL_ROWS:
                    truncated = True
                    break
                cells = ["" if cell is None else str(cell).strip() for cell in row]
                if any(cells):
                    rows_out.append(" | ".join(cells))
            if truncated:
                warnings.append(
                    f"Sheet “{name}” was cut after {MAX_EXCEL_ROWS} rows. Paste or split if the rest matters."
                )
            if rows_out:
                blocks.append(f"## {name}\n" + "\n".join(rows_out))
    finally:
        workbook.close()
    return "\n\n".join(blocks), warnings


def _csv(data: bytes) -> tuple[str, list[str]]:
    text, warnings = _text(data)
    reader = csv.reader(io.StringIO(text))
    rows: list[str] = []
    truncated = False
    for index, row in enumerate(reader, start=1):
        if index > MAX_EXCEL_ROWS:
            truncated = True
            break
        if any(cell.strip() for cell in row):
            rows.append(" | ".join(cell.strip() for cell in row))
    if truncated:
        warnings.append(
            f"Only the first {MAX_EXCEL_ROWS} rows were read. Paste or split if the rest matters."
        )
    return "\n".join(rows), warnings


def _text(data: bytes) -> tuple[str, list[str]]:
    try:
        return data.decode("utf-8"), []
    except UnicodeDecodeError:
        return data.decode("utf-8", errors="replace"), [
            "Some characters could not be read as UTF-8 and were replaced."
        ]
